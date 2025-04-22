from flask import Flask, render_template, request, redirect, url_for, session,jsonify
from flask_mysqldb import MySQL
import MySQLdb.cursors
import re
import pickle
import pandas as pd
import xgboost as xgb
from sklearn.model_selection import train_test_split
import warnings
import numpy as np
import random
from random import choice, sample
import PyPDF2
from docx import Document

model = pickle.load(open("sih1.pkl", "rb"))
  
  
app = Flask(__name__)
  
  
app.secret_key = 'xyzsdfg'
  
app.config['MYSQL_HOST'] = 'sql12.freesqldatabase.com'
app.config['MYSQL_USER'] = 'sql12774742'
app.config['MYSQL_PASSWORD'] = '49LqPLGIZw'
app.config['MYSQL_DB'] = 'sql12774742'
  
mysql = MySQL(app)
  
@app.route('/')
# editing it now :-

# @app.route('/login01a', methods =['GET', 'POST'])
# def login01a():
#     mesage = ''
#     if request.method == 'POST' and 'email' in request.form and 'password' in request.form:
#         email = request.form['email']
#         password = request.form['password']
#         cursor = mysql.connect.cursor(MySQLdb.cursors.DictCursor)
#         cursor.execute('SELECT * FROM users WHERE email = % s AND password = % s', (email, password, ))
#         user = cursor.fetchone()
#         if user:
#             session['loggedin'] = True
#             # session['userid'] = user['userid']
#             session['name'] = user['name']
#             email_a = user['email']
#             # branch= user['Branch']
#             # roll_no= user['Roll_no']
#             # score_10= user['10th_Score']
#             # score_12= user['12th_Score']
#             # cgpa= user['CGPA']
#             # attendence= user['Attendance']
#             # mesage = 'Logged in successfully !'

#             #get data
#             # cursor = mysql.connect.cursor(MySQLdb.cursors.DictCursor)
#             # cursor.execute("SELECT name, email FROM user WHERE Branch='AIML';")
#             # myresult=cursor.fetchall()
#             return render_template('home.html', d=email_a)
#         else:
#             mesage = 'Please enter correct email / password !'
#     return render_template('login.html', mesage = mesage)


# @app.route('/register01', methods =['GET', 'POST'])
# def register01():
#     mesage = ''
#     if request.method == 'POST' and 'name' in request.form and 'password' in request.form and 'email' in request.form :
#         userName = request.form['name']
#         password = request.form['password']
#         email = request.form['email']

#         try:
#             conn = mysql.connect

#             if conn:
#                 cursor = conn.cursor(MySQLdb.cursors.DictCursor)
#         # cursor = mysql.connect.cursor(MySQLdb.cursors.DictCursor)
#                 cursor.execute('SELECT * FROM usersss WHERE email = % s', (email, ))
#                 account = cursor.fetchone()
#                 if account:
#                     mesage = 'Account already exists !'
#                 elif not re.match(r'[^@]+@[^@]+\.[^@]+', email):
#                     mesage = 'Invalid email address !'
#                 elif not userName or not password or not email:
#                     mesage = 'Please fill out the form !'
#                 else:
#                     cursor.execute('INSERT INTO users (name, email, password) VALUES (% s, % s, % s)', (userName, email, password ))
#                     conn.commit()
#                     cursor.close()
#                     mesage = 'You have successfully registered !'

#             else:
#                 return "Connection not established."
            
#         except Exception as e:
#             return f"An error occurred: {str(e)}"
                
#     elif request.method == 'POST':
#         mesage = 'Please fill out the form !'
#     return render_template('register01.html', mesage = mesage)

@app.route('/login', methods =['GET', 'POST'])
def login():
    mesage = ''
    if request.method == 'POST' and 'email' in request.form and 'password' in request.form:
        email = request.form['email']
        password = request.form['password']
        cursor = mysql.connect.cursor(MySQLdb.cursors.DictCursor)
        cursor.execute('SELECT * FROM users WHERE email = % s AND password = % s', (email, password, ))
        # cursor.execute(f'SELECT * FROM user WHERE email = "{email}" AND password = "{password}"')
        user = cursor.fetchone()
        if user:
            session['loggedin'] = True
            # session['userid'] = user['userid']
            session['name'] = user['name']
            email_a = user['email']
            # branch= user['Branch']
            # roll_no= user['Roll_no']
            # score_10= user['10th_Score']
            # score_12= user['12th_Score']
            # cgpa= user['CGPA']
            # attendence= user['Attendance']
            # mesage = 'Logged in successfully !'
            return render_template('home.html', d=email_a)
        else:
            mesage = 'Please enter correct email / password !'
    return render_template('login.html', mesage = mesage)
  
# @app.route('/logout')
# def logout():
#     session.pop('loggedin', None)
#     # session.pop('userid', None)
#     session.pop('email', None)
#     # #yaha se
#     # session.clear()
#     # #yaha tak
#     return redirect(url_for('login'))
  
@app.route('/register', methods =['GET', 'POST'])
def register():
    mesage = ''
    if request.method == 'POST' and 'name' in request.form and 'password' in request.form and 'email' in request.form :
        userName = request.form['name']
        password = request.form['password']
        email = request.form['email']

        try:
            conn = mysql.connect

            if conn:
                cursor = conn.cursor(MySQLdb.cursors.DictCursor)
        # cursor = mysql.connect.cursor(MySQLdb.cursors.DictCursor)
                cursor.execute('SELECT * FROM users WHERE email = % s', (email, ))
                account = cursor.fetchone()
                if account:
                    mesage = 'Account already exists !'
                elif not re.match(r'[^@]+@[^@]+\.[^@]+', email):
                    mesage = 'Invalid email address !'
                elif not userName or not password or not email:
                    mesage = 'Please fill out the form !'
                else:
                    cursor.execute('INSERT INTO users (name, email, password) VALUES (% s, % s, % s)', (userName, email, password ))
                    # cursor.execute(f'INSERT INTO user (name, email, password) VALUES ({userName}, {email}, {password})')
                    conn.commit()
                    cursor.close()
                    mesage = 'You have successfully registered !'

            else:
                return "Connection not established."
            
        except Exception as e:
            return f"An error occurred: {str(e)}"
                
    elif request.method == 'POST':
        mesage = 'Please fill out the form !'
    return render_template('register.html', mesage = mesage)



@app.route('/login01a', methods =['GET', 'POST'])
def login01a():
    mesage = ''
    if request.method == 'POST' and 'email' in request.form and 'password' in request.form:
        email = request.form['email']
        password = request.form['password']
        cursor = mysql.connect.cursor(MySQLdb.cursors.DictCursor)
        cursor.execute('SELECT * FROM users WHERE email = % s AND password = % s', (email, password, ))
        user = cursor.fetchone()
        if user:
            session['loggedin'] = True
            # session['userid'] = user['userid']
            session['name'] = user['name']
            email_a = user['email']
            # branch= user['Branch']
            # roll_no= user['Roll_no']
            # score_10= user['10th_Score']
            # score_12= user['12th_Score']
            # cgpa= user['CGPA']
            # attendence= user['Attendance']
            # mesage = 'Logged in successfully !'

            #get data
            # cursor = mysql.connect.cursor(MySQLdb.cursors.DictCursor)
            # cursor.execute("SELECT name, email FROM user WHERE Branch='AIML';")
            # myresult=cursor.fetchall()
            return render_template('home01.html', d=email_a)
        else:
            mesage = 'Please enter correct email / password !'
    return render_template('login01.html', mesage = mesage)


@app.route('/register01', methods =['GET', 'POST'])
def register01():
    mesage = ''
    if request.method == 'POST' and 'name' in request.form and 'password' in request.form and 'email' in request.form :
        userName = request.form['name']
        password = request.form['password']
        email = request.form['email']

        try:
            conn = mysql.connect

            if conn:
                cursor = conn.cursor(MySQLdb.cursors.DictCursor)
        # cursor = mysql.connect.cursor(MySQLdb.cursors.DictCursor)
                cursor.execute('SELECT * FROM users WHERE email = % s', (email, ))
                account = cursor.fetchone()
                if account:
                    mesage = 'Account already exists !'
                elif not re.match(r'[^@]+@[^@]+\.[^@]+', email):
                    mesage = 'Invalid email address !'
                elif not userName or not password or not email:
                    mesage = 'Please fill out the form !'
                else:
                    cursor.execute('INSERT INTO users (name, email, password) VALUES (% s, % s, % s)', (userName, email, password ))
                    conn.commit()
                    cursor.close()
                    mesage = 'You have successfully registered !'

            else:
                return "Connection not established."
            
        except Exception as e:
            return f"An error occurred: {str(e)}"
                
    elif request.method == 'POST':
        mesage = 'Please fill out the form !'
    return render_template('register01.html', mesage = mesage)


@app.route('/login01')
def login01():
    return render_template('login01.html')



# yaha tak :)


@app.route('/home')
def home():
    return render_template('home.html')

# @app.route('/test', methods =['GET', 'POST'])
# def test():
#     data1=request.files['pdf']

#     # pip install PyPDF2 python-docx nltk spacy scikit-learn
#     def pdf_to_text(pdf_file):
#         with open(pdf_file, 'rb') as file:
#             reader = PyPDF2.PdfReader(file)
#             text = ''
#             for page in reader.pages:
#                 text += page.extract_text()
#         return text
#     from docx import Document
#     def docx_to_text(docx_file):
#         doc = Document(docx_file)
#         return '\n'.join([para.text for para in doc.paragraphs])
#     def extract_skills_section(text):
#         lines = text.splitlines()
#         skills_section = ""
#         capture = False

#         for line in lines:
#             if "skills" in line.lower():
#                 capture = True
#             elif capture:
#                 if line.strip() == "" or any(header in line.lower() for header in ["experience", "education", "work"]):
#                     break
#                 skills_section += line + " "

#         return skills_section.strip()
#     def extract_skills(skills_section):
#         skills = skills_section.split(",")  # Assume skills are comma-separated
#         return [skill.strip() for skill in skills if skill.strip()]
    
#     def process_resume(file):
#         if file.endswith('.pdf'):
#             text = pdf_to_text(file)
#         elif file.endswith('.docx'):
#             text = docx_to_text(file)
#         else:
#             raise ValueError("Unsupported file type")

#         skills_section = extract_skills_section(text)
#         skills = extract_skills(skills_section)
#         skills_string = ' '.join(skills)
# # pr    int(skills_string)

#         return skills_string
    
#     skills=process_resume(data1)
#     return render_template('after.html', s=skills)

@app.route('/test01', methods=['POST'])
def test01():
    if 'pdf' not in request.files:
        # flash('No file part in the request')
        return redirect(url_for('upload_page'))  # Redirect to an upload page

    file = request.files['pdf']
    pos=request.form['position']
    try:
        conn = mysql.connect  # Assuming you have mysql configured
        if conn:
            cursor = conn.cursor(MySQLdb.cursors.DictCursor)
            query = "SELECT skills FROM openings WHERE position = %s"
            cursor.execute(query, (pos,))
            results = cursor.fetchone()
            # results = cursor.fetchall()
            cursor.close()
            conn.commit()
            if results and 'skills' in results:
                sk = results['skills']
            else:
                sk = ""  # Handle the case where no matching position is found
        else:
            return "Connection not established"
    except Exception as e:
            return f"An error occurred: {str(e)}"
    
    if file.filename == '':
        # flash('No selected file')
        return redirect(url_for('upload_page'))

    if file:
        # Get the file extension
        file_extension = file.filename.rsplit('.', 1)[1].lower()

        def pdf_to_text(pdf_file):
            reader = PyPDF2.PdfReader(pdf_file)
            text = ''
            for page in reader.pages:
                text += page.extract_text()
            return text

        def docx_to_text(docx_file):
            doc = Document(docx_file)
            return '\n'.join([para.text for para in doc.paragraphs])

        def extract_skills_section(text):
            lines = text.splitlines()
            skills_section = ""
            capture = False

            for line in lines:
                if "skills" in line.lower():
                    capture = True
                elif capture:
                    if line.strip() == "" or any(header in line.lower() for header in ["experience", "education", "work"]):
                        break
                    skills_section += line + " "

            return skills_section.strip()

        # def extract_skills_section(text):
        #     lines = text.splitlines()
        #     capture = False
        #     skills_lines = []
        #     section_headers = ["experience", "education", "project", "certifications", "summary", "work", "internship", "achievement", "objective"]
        #     max_skills_lines = 10
        
        #     for i, line in enumerate(lines):
        #         stripped_line = line.strip()
        #         lower_line = stripped_line.lower()
        
        #         if "skill" in lower_line and len(lower_line) <= 30:
        #             capture = True
        #             continue
                
        #         if capture:
        #             if any(header in lower_line for header in section_headers) or len(skills_lines) >= max_skills_lines:
        #                 break
        #             if 2 < len(stripped_line.split()) < 20:
        #                 skills_lines.append(stripped_line)
        
        #     extracted = ' '.join(skills_lines)
        #     extracted = re.sub(r'[^a-zA-Z0-9,+.-]', ' ', extracted)
        
        #     parts = re.split(r',|•|-|•', extracted)
        #     skill_candidates = []
        #     for part in parts:
        #         clean = part.strip()
        #         if 1 < len(clean.split()) <= 5:
        #             skill_candidates.append(clean)
        
        #     final_skills = ', '.join(sorted(set(skill_candidates), key=skill_candidates.index))
        #     return final_skills

        def extract_skills(skills_section):
            skills = skills_section.split(",")  # Assume skills are comma-separated
            return [skill.strip() for skill in skills if skill.strip()]

        def process_resume(file):
            if file_extension == 'pdf':
                text = pdf_to_text(file)
            elif file_extension == 'docx':
                text = docx_to_text(file)
            else:
                raise ValueError("Unsupported file type")

            skills_section = extract_skills_section(text)
            skills = extract_skills(skills_section)
            skills_string = ', '.join(skills)

            return skills_string
        
        skills = process_resume(file)
        # def normalize_skills_string(skills_string):
            # Lowercase, split by comma, strip each skill
            # return [skill.strip().lower() for skill in skills_string.split(',') if skill.strip()]

        def normalize_skills_string(skills_string):
            # Replace anything that's not a letter or + with space, then split
            cleaned = re.sub(r'[^a-zA-Z+]', ' ', skills_string.lower())
            return [skill.strip() for skill in cleaned.split() if skill.strip()]


        # def replace_non_letters_with_comma(input_string):
        #     # Replace all characters that are not letters (a-z or A-Z) with a comma
        #     return re.sub(r'[^a-zA-Z+]', ',', input_string)
        # skills=replace_non_letters_with_comma(skills)
        candidate_skills_list = normalize_skills_string(skills)
        position_skills_list = normalize_skills_string(sk)
        print("Raw candidate resume skills string:", skills)
        print("Raw position DB skills string:", sk)

        

        

        # def find_common_skills(text1, text2):
        #     # Split the texts by commas and convert to sets
        #     skills1 = set(text1.split(','))
        #     skills2 = set(text2.split(','))

        #     # Find common skills
        #     common_skills = skills1.intersection(skills2)

        #     # Join the common skills into a string separated by commas
        #     return ','.join(common_skills)
        def find_common_and_unmatched_skills(candidate_skills_list, position_skills_list):
            candidate_skills = set(candidate_skills_list)
            position_skills = set(position_skills_list)
        
            matched_skills = candidate_skills & position_skills
            unmatched_skills = position_skills - candidate_skills
        
            return ', '.join(sorted(matched_skills)), ', '.join(sorted(unmatched_skills))
        
        candidate_skills_list = normalize_skills_string(skills)
        position_skills_list = normalize_skills_string(sk)

        matched_skills, unmatched_skills = find_common_and_unmatched_skills(candidate_skills_list, position_skills_list)

        return render_template(
            'home01.html',
            matched=matched_skills,
            unmatched=unmatched_skills,
            s1=candidate_skills_list,
            s2=position_skills_list,
            candidate=', '.join(candidate_skills_list),
            position=', '.join(position_skills_list)
        )

        
        # matched_skills, unmatched_skills = find_common_and_unmatched_skills(skills, sk)
        # return render_template('home01.html', matched=matched_skills, unmatched=unmatched_skills, candidate=skills, position=sk)

        # skills1=find_common_skills(skills,sk)
        # return render_template('home01.html', s=skills1,s1=skills,s2=sk)

@app.route('/test', methods=['POST'])
def test():
    if 'pdf' not in request.files:
        # flash('No file part in the request')
        return redirect(url_for('upload_page'))  # Redirect to an upload page

    file = request.files['pdf']
    pos=request.form['position']
    try:
        conn = mysql.connect  # Assuming you have mysql configured
        if conn:
            cursor = conn.cursor(MySQLdb.cursors.DictCursor)
            query = "SELECT skills FROM openings WHERE position = %s"
            cursor.execute(query, (pos,))
            results = cursor.fetchone()
            # results = cursor.fetchall()
            cursor.close()
            conn.commit()
            if results and 'skills' in results:
                sk = results['skills']
            else:
                sk = ""  # Handle the case where no matching position is found
        else:
            return "Connection not established"
    except Exception as e:
            return f"An error occurred: {str(e)}"
    
    if file.filename == '':
        # flash('No selected file')
        return redirect(url_for('upload_page'))

    if file:
        # Get the file extension
        file_extension = file.filename.rsplit('.', 1)[1].lower()

        def pdf_to_text(pdf_file):
            reader = PyPDF2.PdfReader(pdf_file)
            text = ''
            for page in reader.pages:
                text += page.extract_text()
            return text

        def docx_to_text(docx_file):
            doc = Document(docx_file)
            return '\n'.join([para.text for para in doc.paragraphs])

        def extract_skills_section(text):
            lines = text.splitlines()
            skills_section = ""
            capture = False

            for line in lines:
                if "skills" in line.lower():
                    capture = True
                elif capture:
                    if line.strip() == "" or any(header in line.lower() for header in ["experience", "education", "work"]):
                        break
                    skills_section += line + " "

            return skills_section.strip()

        # def extract_skills_section(text):
        #     lines = text.splitlines()
        #     capture = False
        #     skills_lines = []
        #     section_headers = ["experience", "education", "project", "certifications", "summary", "work", "internship", "achievement", "objective"]
        #     max_skills_lines = 10
        
        #     for i, line in enumerate(lines):
        #         stripped_line = line.strip()
        #         lower_line = stripped_line.lower()
        
        #         if "skill" in lower_line and len(lower_line) <= 30:
        #             capture = True
        #             continue
                
        #         if capture:
        #             if any(header in lower_line for header in section_headers) or len(skills_lines) >= max_skills_lines:
        #                 break
        #             if 2 < len(stripped_line.split()) < 20:
        #                 skills_lines.append(stripped_line)
        
        #     extracted = ' '.join(skills_lines)
        #     extracted = re.sub(r'[^a-zA-Z0-9,+.-]', ' ', extracted)
        
        #     parts = re.split(r',|•|-|•', extracted)
        #     skill_candidates = []
        #     for part in parts:
        #         clean = part.strip()
        #         if 1 < len(clean.split()) <= 5:
        #             skill_candidates.append(clean)
        
        #     final_skills = ', '.join(sorted(set(skill_candidates), key=skill_candidates.index))
        #     return final_skills

        def extract_skills(skills_section):
            skills = skills_section.split(",")  # Assume skills are comma-separated
            return [skill.strip() for skill in skills if skill.strip()]

        def process_resume(file):
            if file_extension == 'pdf':
                text = pdf_to_text(file)
            elif file_extension == 'docx':
                text = docx_to_text(file)
            else:
                raise ValueError("Unsupported file type")

            skills_section = extract_skills_section(text)
            skills = extract_skills(skills_section)
            skills_string = ' '.join(skills)
            return skills_string
        
        skills = process_resume(file)
        def replace_non_letters_with_comma(input_string):
            # Replace all characters that are not letters (a-z or A-Z) with a comma
            return re.sub(r'[^a-zA-Z+]', ',', input_string)
        skills=replace_non_letters_with_comma(skills)
        def find_common_skills(text1, text2):
            # Split the texts by commas and convert to sets
            skills1 = set(text1.split(','))
            skills2 = set(text2.split(','))

            # Find common skills
            common_skills = skills1.intersection(skills2)

            # Join the common skills into a string separated by commas
            return ','.join(common_skills)
        skills1=find_common_skills(skills,sk)
        return render_template('home.html', s=skills1,s1=skills,s2=sk)
    
@app.route('/match_skills', methods=['POST'])
# def match_skills():
#     def get_matching_professions(input_skills):
#         input_skills_set = set(input_skills.lower().split(","))
        
#         try:
#             conn = mysql.connect  # Assuming you have mysql configured
#             if conn:
#                 cursor = conn.cursor(MySQLdb.cursors.DictCursor)
#                 cursor.execute("SELECT professor, skills FROM prof")
#                 results = cursor.fetchall()
#                 cursor.close()
#                 conn.commit()
#             else:
#                 return "Connection not established"

#         except Exception as e:
#             return f"An error occurred: {str(e)}"

#         matches = []
#         for row in results:
#             profession = row['professor']
#             skills = row['skills']
#             profession_skills_set = set(skills.lower().split(","))
#             match_score = len(input_skills_set.intersection(profession_skills_set))
#             if match_score > 0:
#                 matches.append((profession, match_score))

#         matches.sort(key=lambda x: x[1], reverse=True)  # Sort by match score
        
#         # Convert the top 5 matches into a concatenated string
#         top_matches = " ".join([f"{profession}({score})" for profession, score in matches[:5]])
        
#         return top_matches

#     # def get_matching_professions(input_skills):
#     #     input_skills_set = set(input_skills.lower().split(","))
#     #     # conn = sqlite3.connect('database.db')  # Connect to your database
#     #     # cursor = conn.cursor()

#     #     # cursor.execute("SELECT profession_name, skills FROM prof")
#     #     # results = cursor.fetchall()
#     #     try:
#     #         conn = mysql.connect
#     #         if conn:
#     #                 cursor = conn.cursor(MySQLdb.cursors.DictCursor)
#     #                 # cursor = mysql.connect.cursor(MySQLdb.cursors.DictCursor)
#     #                 cursor.execute("SELECT profession_name, skills FROM prof")
#     #                 results = cursor.fetchall()
#     #                 conn.commit()
#     #                 cursor.close()
#     #                 # return render_template("success.html")
#     #                 # message="success"
#     #         else:
#     #             return "connection not extablished"
    
#     #     except Exception as e:
#     #         return f"An error occurred: {str(e)}"

#     #     matches = []
#     #     for profession, skills in results:
#     #         profession_skills_set = set(skills.lower().split(","))
#     #         match_score = len(input_skills_set.intersection(profession_skills_set))
#     #         if match_score > 0:
#     #             matches.append((profession, match_score))

#     #     matches.sort(key=lambda x: x[1], reverse=True)  # Sort by match score
#     #     return matches[:5]  # Return top 5

#     input_skills = request.form['skills']
#     top_matches = get_matching_professions(input_skills)
#     return render_template('after1.html', s=top_matches)
# def match_skills():
#     def get_matching_professions(input_skills):
#         input_skills_set = set(input_skills.lower().split(","))

#         try:
#             conn = mysql.connect  # Assuming you have mysql configured
#             if conn:
#                 cursor = conn.cursor(MySQLdb.cursors.DictCursor)
#                 cursor.execute("SELECT professor, skills FROM prof")
#                 results = cursor.fetchall()
#                 cursor.close()
#                 conn.commit()
#             else:
#                 return "Connection not established"
        
#         except Exception as e:
#             return f"An error occurred: {str(e)}"

#         matches = []
#         for row in results:
#             profession = row['professor']
#             skills = row['skills']
#             profession_skills_set = set(skills.lower().split(","))
#             matched_skills = input_skills_set.intersection(profession_skills_set)
#             match_score = len(matched_skills)
#             if match_score > 0:
#                 matches.append((profession, match_score, matched_skills))

#         # Sort matches by score in descending order
#         matches.sort(key=lambda x: x[1], reverse=True)
        
#         # Format the top 5 matches
#         top_matches = " ".join([f"{profession}({score}) - Matched skills: {', '.join(matched_skills)}" 
#                                 for profession, score, matched_skills in matches[:5]])
        
#         return top_matches

#     input_skills = request.form['skills']
#     top_matches = get_matching_professions(input_skills)
#     return render_template('after1.html', s=top_matches)
def match_skills():
    def get_matching_professions(input_skills):
        input_skills_set = set(input_skills.lower().split(","))

        try:
            conn = mysql.connect  # Assuming you have mysql configured
            if conn:
                cursor = conn.cursor(MySQLdb.cursors.DictCursor)
                cursor.execute("SELECT professor, skills FROM prof")
                results = cursor.fetchall()
                cursor.close()
                conn.commit()
            else:
                return "Connection not established"
        
        except Exception as e:
            return f"An error occurred: {str(e)}"

        matches = []
        for row in results:
            profession = row['professor']
            skills = row['skills']
            profession_skills_set = set(skills.lower().split(","))
            matched_skills = input_skills_set.intersection(profession_skills_set)
            match_score = len(matched_skills)
            if match_score > 0:
                matches.append((profession, match_score, matched_skills))

        # Sort matches by score in descending order
        matches.sort(key=lambda x: x[1], reverse=True)
        
        # Format the top 5 matches
        top_matches = "<br>".join([f"{profession} - Score: {score} - Matched skills: {', '.join(matched_skills)}" 
                                   for profession, score, matched_skills in matches[:5]])
        
        return top_matches

    input_skills = request.form['skills']
    top_matches = get_matching_professions(input_skills)
    return render_template('after1.html', s=top_matches)








if __name__ == "__main__":
    app.run('0.0.0.0',port=8080,debug=True)
